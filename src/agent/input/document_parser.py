"""
Document Parser for SRS Documents.

This module provides parsers for different document formats (DOCX, PDF, text)
to extract requirements from SRS documents.
"""

import logging
import re
import os
from typing import Dict, List, Optional, Any

# Setup logging
logger = logging.getLogger(__name__)

class DocumentParser:
    """
    Base class for document parsers.
    
    This class defines the interface for document parsers and provides
    common functionality for requirement extraction.
    """
    
    def __init__(self):
        """Initialize the document parser."""
        logger.debug("Initializing document parser")
    
    def parse(self, file_path: str) -> List[Dict]:
        """
        Parse a document and extract requirements.
        
        Args:
            file_path: Path to the document
            
        Returns:
            List of requirement dictionaries
        """
        raise NotImplementedError("Subclasses must implement parse()")
    
    def _extract_requirements(self, text: str) -> List[Dict]:
        """
        Extract requirements from text using regex patterns.
        
        Args:
            text: The text to extract requirements from
            
        Returns:
            List of requirement dictionaries
        """
        logger.debug("Extracting requirements from text")
        requirements = []
        
        # Pattern for requirement IDs (e.g., REQ-001, R-123, etc.)
        id_pattern = r'(?:REQ|R)-\d+'
        
        # Split text into sections that might contain requirements
        sections = re.split(r'\n\s*\n', text)
        
        for section in sections:
            # Try to find a requirement ID
            id_match = re.search(id_pattern, section)
            req_id = id_match.group(0) if id_match else None
            
            # If no ID found, try to determine if this is a requirement by keywords
            is_requirement = False
            if not req_id:
                requirement_keywords = ['shall', 'must', 'should', 'will', 'requires']
                if any(keyword in section.lower() for keyword in requirement_keywords):
                    is_requirement = True
                    req_id = f"REQ-AUTO-{len(requirements) + 1:03d}"
            
            # If we have an ID or determined it's a requirement, extract it
            if req_id or is_requirement:
                # Extract description (everything after the ID or the whole section)
                description = section
                if req_id and id_match:
                    description = section[id_match.end():].strip()
                
                # Create requirement dictionary
                requirement = {
                    'id': req_id,
                    'description': description,
                    'source': 'SRS',
                    'tags': self._extract_tags(description)
                }
                
                requirements.append(requirement)
        
        logger.info(f"Extracted {len(requirements)} requirements")
        return requirements
    
    def _extract_tags(self, text: str) -> List[str]:
        """
        Extract tags/keywords from requirement text.
        
        Args:
            text: The requirement text
            
        Returns:
            List of tags
        """
        tags = []
        
        # Check for criticality indicators
        if any(word in text.lower() for word in ['critical', 'essential', 'mandatory']):
            tags.append('critical')
        elif any(word in text.lower() for word in ['important', 'significant']):
            tags.append('important')
        
        # Check for functional areas
        if 'user interface' in text.lower() or 'ui' in text.lower():
            tags.append('ui')
        if 'database' in text.lower() or 'data' in text.lower():
            tags.append('data')
        if 'security' in text.lower() or 'authentication' in text.lower():
            tags.append('security')
        
        return tags


class DocxParser(DocumentParser):
    """Parser for Microsoft Word (.docx) documents."""
    
    def parse(self, file_path: str) -> List[Dict]:
        """
        Parse a .docx document and extract requirements.
        
        Args:
            file_path: Path to the .docx document
            
        Returns:
            List of requirement dictionaries
        """
        logger.info(f"Parsing DOCX document: {file_path}")
        
        try:
            # Import here to avoid dependency if not needed
            from docx import Document
            
            document = Document(file_path)
            full_text = []
            
            # Extract text from paragraphs
            for para in document.paragraphs:
                full_text.append(para.text)
            
            # Extract text from tables
            for table in document.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text)
                    full_text.append(' | '.join(row_text))
            
            # Join all text with newlines
            document_text = '\n'.join(full_text)
            
            # Extract requirements from the text
            return self._extract_requirements(document_text)
            
        except ImportError:
            logger.error("python-docx package not installed. Install with: pip install python-docx")
            raise
        except Exception as e:
            logger.error(f"Error parsing DOCX document: {str(e)}")
            raise


class PdfParser(DocumentParser):
    """Parser for PDF documents."""
    
    def parse(self, file_path: str) -> List[Dict]:
        """
        Parse a PDF document and extract requirements.
        
        Args:
            file_path: Path to the PDF document
            
        Returns:
            List of requirement dictionaries
        """
        logger.info(f"Parsing PDF document: {file_path}")
        
        try:
            # Import here to avoid dependency if not needed
            import PyPDF2
            
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                full_text = []
                
                # Extract text from each page
                for page_num in range(len(reader.pages)):
                    page = reader.pages[page_num]
                    full_text.append(page.extract_text())
                
                # Join all text with newlines
                document_text = '\n'.join(full_text)
                
                # Extract requirements from the text
                return self._extract_requirements(document_text)
                
        except ImportError:
            logger.error("PyPDF2 package not installed. Install with: pip install PyPDF2")
            raise
        except Exception as e:
            logger.error(f"Error parsing PDF document: {str(e)}")
            raise


class TextParser(DocumentParser):
    """Parser for plain text documents."""
    
    def parse(self, file_path: str) -> List[Dict]:
        """
        Parse a text document and extract requirements.
        
        Args:
            file_path: Path to the text document
            
        Returns:
            List of requirement dictionaries
        """
        logger.info(f"Parsing text document: {file_path}")
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                document_text = file.read()
                
                # Extract requirements from the text
                return self._extract_requirements(document_text)
                
        except Exception as e:
            logger.error(f"Error parsing text document: {str(e)}")
            raise


def get_parser_for_file(file_path: str) -> DocumentParser:
    """
    Get the appropriate parser for a file based on its extension.
    
    Args:
        file_path: Path to the document
        
    Returns:
        An instance of the appropriate DocumentParser subclass
    """
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    
    if ext == '.docx':
        return DocxParser()
    elif ext == '.pdf':
        return PdfParser()
    else:
        return TextParser()
